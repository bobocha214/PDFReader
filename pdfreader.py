# # -*- coding: utf-8 -*-
# # @Time : 2022/7/19 11:45
# # @Author : Shengzhiyu
# # coding:utf-8
import jieba
import re
import docx
# import PyPDF2
import pdfplumber
import os
from PIL import Image
import numpy as np
from wordcloud import WordCloud
from collections import Counter
'''传入pdf文件地址，读取pdf文本，返回读出的文本字符串'''
filePath="pdf" # 文件夹路径
fileList=os.listdir(filePath)
jieba.load_userdict('./userdict.txt')
finalwords= {}

stopword = [line.rstrip() for line in open("stopwords.txt", 'r', encoding='utf-8')]
def extract_content():
    all=''
    for file in fileList:
        path = os.path.join(filePath, file)
        print(file)  # 文件名
        with pdfplumber.open(path) as pdf_file:
            content = ''
            for i in range(len(pdf_file.pages)):
                # print("当前第 %s 页" % i)
                page_text = pdf_file.pages[i]
                page_content = page_text.extract_text()
                if page_content:
                    content = content + page_content
                    # print(page_content)
        all+=content

    return all


# 文本清洗
def clean_text(text):
    newtext = []
    text = re.sub(r'[A-Za-z0-9\s+\.\!\/_,$%^*：()（）?;；，．:-【】+\"\']+|[+——、~@#￥%……&*()<>-]+', '', text)  # 去除数字
    text = re.findall(r'[\u4e00-\u9fa5a-zA-Z0-9]',text)
    text = jieba.lcut(''.join(text))  # 分词
    for word in text:
        if word not in stopword:  # 去停用词 + 词性筛选
            newtext.append(word)
    lineswords = ' '.join(newtext)
    return lineswords


# 统计词频
def counter_word(data):
    lines = ''
    for line in data:
        lines += line
    data_list = lines.split(' ')
    words_count = Counter(data_list)
    # print(words_count)
    count_res = words_count.most_common()
    return count_res

if __name__ == '__main__':
    result = extract_content()
    clean_text=clean_text(result)
    counter_result=counter_word(clean_text)
    stopwords = ['我', '和', '你', '的', '地', '得', '了', '都', '对', '向', '在', '可', '能', '为', '要', '再', '是', '等', '一', '二',
                 '三', '四', '五', '六', '七', '八', '九', '十', '各位', '代表', '一年', '请予']
    #字体
    font = r'C:\Windows\Fonts\simhei.ttf'
    #词云背景形状
    background = Image.open("img_new.png").convert('RGB')
    graph = np.array(background)
    wordcloud = WordCloud(background_color='white', scale=2, max_words=500, max_font_size=150,font_path=font,
                          stopwords=stopwords, mask=graph, colormap='brg').generate_from_text(clean_text)

    wordcloud.to_file("alice.png")
    print(counter_result)
    finalwords=counter_result

# 存入result.docx中
doc = docx.Document()
doc.add_paragraph('词频：')
for i in finalwords:
    doc.add_paragraph("{}".format(i))
doc.save('result.docx')